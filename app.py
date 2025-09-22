import streamlit as st
from openai import OpenAI
from io import BytesIO
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from docx.shared import Inches
import requests
import time

# Page configuration
st.set_page_config(
    page_title="✨ Ebook Writer", 
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# Custom CSS for professional styling
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    
    .main-container {
        font-family: 'Inter', sans-serif;
    }
    
    .hero-section {
        background: linear-gradient(135deg, #1e3a8a 0%, #3b82f6 50%, #8b5cf6 100%);
        margin: -1rem -1rem 0 -1rem;
        padding: 4rem 2rem;
        position: relative;
        overflow: hidden;
    }
    
    .hero-section::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        bottom: 0;
        background: url("data:image/svg+xml,%3Csvg width='60' height='60' viewBox='0 0 60 60' xmlns='http://www.w3.org/2000/svg'%3E%3Cg fill='none' fill-rule='evenodd'%3E%3Cg fill='%23ffffff' fill-opacity='0.05'%3E%3Ccircle cx='30' cy='30' r='2'/%3E%3C/g%3E%3C/g%3E%3C/svg%3E");
    }
    
    .hero-content {
        position: relative;
        z-index: 2;
        text-align: center;
        max-width: 800px;
        margin: 0 auto;
    }
    
    .hero-badge {
        display: inline-block;
        background: rgba(255, 255, 255, 0.2);
        color: white;
        padding: 0.5rem 1.5rem;
        border-radius: 50px;
        font-size: 0.9rem;
        font-weight: 500;
        margin-bottom: 1.5rem;
        backdrop-filter: blur(10px);
        border: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .hero-title {
        color: white;
        font-size: 3.5rem;
        font-weight: 700;
        margin-bottom: 1rem;
        text-shadow: 0 4px 20px rgba(0,0,0,0.3);
        line-height: 1.2;
    }
    
    .hero-subtitle {
        color: rgba(255,255,255,0.9);
        font-size: 1.3rem;
        font-weight: 400;
        margin-bottom: 2rem;
        line-height: 1.6;
    }
    
    .hero-stats {
        display: flex;
        justify-content: center;
        gap: 3rem;
        margin-top: 2rem;
    }
    
    .stat-item {
        text-align: center;
        color: white;
    }
    
    .stat-number {
        font-size: 2rem;
        font-weight: 700;
        display: block;
    }
    
    .stat-label {
        font-size: 0.9rem;
        opacity: 0.8;
    }
    
    .features-section {
        background: #f8fafc;
        margin: 0 -1rem;
        padding: 4rem 2rem;
    }
    
    .features-container {
        max-width: 1200px;
        margin: 0 auto;
    }
    
    .section-title {
        text-align: center;
        font-size: 2.5rem;
        font-weight: 700;
        color: #1e293b;
        margin-bottom: 1rem;
    }
    
    .section-subtitle {
        text-align: center;
        font-size: 1.1rem;
        color: #64748b;
        margin-bottom: 3rem;
        max-width: 600px;
        margin-left: auto;
        margin-right: auto;
    }
    
    .features-grid {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(300px, 1fr));
        gap: 2rem;
        margin-bottom: 3rem;
    }
    
    .feature-card {
        background: white;
        padding: 2rem;
        border-radius: 20px;
        box-shadow: 0 10px 40px rgba(0,0,0,0.08);
        border: 1px solid #e2e8f0;
        transition: all 0.3s ease;
        position: relative;
        overflow: hidden;
    }
    
    .feature-card:hover {
        transform: translateY(-5px);
        box-shadow: 0 20px 60px rgba(0,0,0,0.12);
    }
    
    .feature-icon {
        width: 60px;
        height: 60px;
        background: linear-gradient(135deg, #3b82f6, #8b5cf6);
        border-radius: 16px;
        display: flex;
        align-items: center;
        justify-content: center;
        font-size: 1.5rem;
        margin-bottom: 1.5rem;
    }
    
    .feature-title {
        font-size: 1.3rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 0.8rem;
    }
    
    .feature-description {
        color: #64748b;
        line-height: 1.6;
        font-size: 0.95rem;
    }
    
    .process-section {
        background: white;
        margin: 0 -1rem;
        padding: 4rem 2rem;
    }
    
    .process-container {
        max-width: 900px;
        margin: 0 auto;
    }
    
    .process-steps {
        display: grid;
        grid-template-columns: repeat(auto-fit, minmax(250px, 1fr));
        gap: 2rem;
    }
    
    .process-step {
        text-align: center;
        position: relative;
    }
    
    .step-number {
        width: 50px;
        height: 50px;
        background: linear-gradient(135deg, #3b82f6, #8b5cf6);
        color: white;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        font-weight: 600;
        font-size: 1.1rem;
        margin: 0 auto 1rem;
    }
    
    .step-title {
        font-size: 1.1rem;
        font-weight: 600;
        color: #1e293b;
        margin-bottom: 0.5rem;
    }
    
    .step-description {
        color: #64748b;
        font-size: 0.9rem;
        line-height: 1.5;
    }
    
    .cta-section {
        background: linear-gradient(135deg, #1e293b 0%, #334155 100%);
        margin: 0 -1rem;
        padding: 4rem 2rem;
        text-align: center;
    }
    
    .cta-title {
        color: white;
        font-size: 2.2rem;
        font-weight: 700;
        margin-bottom: 1rem;
    }
    
    .cta-subtitle {
        color: rgba(255, 255, 255, 0.8);
        font-size: 1.1rem;
        margin-bottom: 2rem;
        max-width: 500px;
        margin-left: auto;
        margin-right: auto;
    }
    
    .scene-container {
        background: #f8f9ff;
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        border: 2px solid #e1e8ff;
    }
    
    .download-section {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        margin: 2rem 0;
    }
    
    .download-section h3 {
        color: white;
        margin-bottom: 1rem;
    }
    
    .api-key-container {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
    }
    
    .success-message {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(102, 126, 234, 0.3);
    }
    
    .stTextArea textarea {
        border-radius: 10px;
        border: 2px solid #e1e8ff;
    }
    
    
    .scene-container {
        background: #f8f9ff;
        border-radius: 15px;
        padding: 1.5rem;
        margin: 1rem 0;
        border: 2px solid #e1e8ff;
    }
    
    .download-section {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
        padding: 2rem;
        border-radius: 15px;
        text-align: center;
        margin: 2rem 0;
    }
    
    .download-section h3 {
        color: white;
        margin-bottom: 1rem;
    }
    
    .api-key-container {
        background: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 10px;
        padding: 1rem;
        margin: 1rem 0;
    }
    
    .success-message {
        background: #d4edda;
        border: 1px solid #c3e6cb;
        color: #155724;
        padding: 1rem;
        border-radius: 10px;
        margin: 1rem 0;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #3b82f6 0%, #8b5cf6 100%);
        color: white;
        border: none;
        border-radius: 25px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(59, 130, 246, 0.3);
    }
    
    .stTextArea textarea {
        border-radius: 10px;
        border: 2px solid #e1e8ff;
    }
    
    .stTextInput input {
        border-radius: 10px;
        border: 2px solid #e1e8ff;
    }
    
    @media (max-width: 768px) {
        .hero-title {
            font-size: 2.5rem;
        }
        
        .hero-stats {
            flex-direction: column;
            gap: 1.5rem;
        }
        
        .features-grid {
            grid-template-columns: 1fr;
        }
        
        .process-steps {
            grid-template-columns: 1fr;
        }
    }
</style>
""", unsafe_allow_html=True)

# Professional Header
st.markdown("""
<div class="hero-section">
    <div class="hero-content">
        <div class="hero-badge">✨ Powered by GPT-4 & DALL-E 3</div>
        <h1 class="hero-title">Transform Ideas Into<br>Professional Ebooks</h1>
        <p class="hero-subtitle">
            Create compelling narratives with AI-generated scenes and stunning illustrations. 
            From concept to publication-ready content in minutes.
        </p>
        <div class="hero-stats">
            <div class="stat-item">
                <span class="stat-number">10x</span>
                <span class="stat-label">Faster Writing</span>
            </div>
            <div class="stat-item">
                <span class="stat-number">AI</span>
                <span class="stat-label">Powered Creativity</span>
            </div>
            <div class="stat-item">
                <span class="stat-number">∞</span>
                <span class="stat-label">Story Possibilities</span>
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

# API Key Protection
def check_api_key():
    if "api_key_validated" not in st.session_state:
        st.session_state["api_key_validated"] = False
    
    if not st.session_state["api_key_validated"]:
        st.markdown("""
        <div class="api-key-container">
            <h3>🔐 API Key Required</h3>
            <p>Please enter your OpenAI API key to continue. Your key is stored securely in this session only.</p>
        </div>
        """, unsafe_allow_html=True)
        
        api_key = st.text_input("OpenAI API Key", type="password", placeholder="sk-...")
        
        if st.button("Validate API Key"):
            if api_key.startswith("sk-") and len(api_key) > 20:
                try:
                    # Test the API key
                    test_client = OpenAI(api_key=api_key)
                    test_client.models.list()  # Simple API call to validate
                    
                    st.session_state["openai_api_key"] = api_key
                    st.session_state["api_key_validated"] = True
                    st.markdown("""
                    <div class="success-message">
                        ✅ API Key validated successfully! You can now use the Ebook Writer.
                    </div>
                    """, unsafe_allow_html=True)
                    st.rerun()
                except Exception as e:
                    st.error("❌ Invalid API key. Please check and try again.")
            else:
                st.error("❌ Please enter a valid OpenAI API key (starts with 'sk-')")
        
        return False
    return True

# Main app logic
if check_api_key():
    # Initialize OpenAI client
    openai_client = OpenAI(api_key=st.session_state["openai_api_key"])
    
    # Session state initialization
    if "messages" not in st.session_state:
        st.session_state["messages"] = []
    if "scenes" not in st.session_state:
        st.session_state["scenes"] = []
    if "ebook_text" not in st.session_state:
        st.session_state["ebook_text"] = ""
    if "images" not in st.session_state:
        st.session_state["images"] = {}
    
    # Features Section
    st.markdown("""
    <div class="features-section">
        <div class="features-container">
            <h2 class="section-title">Everything You Need to Create</h2>
            <p class="section-subtitle">
                Professional-grade ebook creation tools powered by cutting-edge AI technology
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Features Grid using Streamlit columns
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🎯</div>
            <div class="feature-title">Smart Story Structure</div>
            <div class="feature-description">
                AI analyzes your concept and creates compelling outlines with proper narrative flow, character development, and engaging plot points.
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">📚</div>
            <div class="feature-title">Export Ready Files</div>
            <div class="feature-description">
                Download publication-ready DOCX and PDF files with embedded images, perfect for self-publishing or sharing.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">✍️</div>
            <div class="feature-title">Rich Scene Generation</div>
            <div class="feature-description">
                Generate vivid, immersive scenes with detailed descriptions, dialogue, and atmospheric elements that bring your story to life.
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">⚡</div>
            <div class="feature-title">Lightning Fast</div>
            <div class="feature-description">
                What traditionally takes weeks can now be accomplished in minutes. From initial concept to finished ebook instantly.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🎨</div>
            <div class="feature-title">AI Illustrations</div>
            <div class="feature-description">
                Create stunning, professional-quality artwork for each scene using DALL-E 3's advanced image generation capabilities.
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div class="feature-card">
            <div class="feature-icon">🔒</div>
            <div class="feature-title">Secure & Private</div>
            <div class="feature-description">
                Your API keys and content remain completely private. No data storage, no tracking, complete creative freedom.
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    # Process Section
    st.markdown("""
    <div class="process-section">
        <div class="process-container">
            <h2 class="section-title">How It Works</h2>
            <p class="section-subtitle">
                Simple 4-step process to transform your ideas into professional ebooks
            </p>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Process Steps using Streamlit columns
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.markdown("""
        <div class="process-step">
            <div class="step-number">1</div>
            <div class="step-title">Share Your Idea</div>
            <div class="step-description">
                Describe your story concept, theme, or inspiration in just a few sentences
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="process-step">
            <div class="step-number">2</div>
            <div class="step-title">Generate Outline</div>
            <div class="step-description">
                AI creates a structured outline with chapters, plot points, and narrative flow
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="process-step">
            <div class="step-number">3</div>
            <div class="step-title">Create Scenes</div>
            <div class="step-description">
                Develop rich, detailed scenes with vivid descriptions and compelling narratives
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown("""
        <div class="process-step">
            <div class="step-number">4</div>
            <div class="step-title">Add Illustrations</div>
            <div class="step-description">
                Generate stunning AI artwork and export your complete ebook
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    # CTA Section
    st.markdown("""
    <div class="cta-section">
        <h2 class="cta-title">Ready to Create Your Masterpiece?</h2>
        <p class="cta-subtitle">
            Join thousands of writers who have already transformed their ideas into professional ebooks
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Helper functions
    def generate_text(prompt):
        with st.spinner("🤖 Generating your ebook outline..."):
            res = openai_client.chat.completions.create(
                model="gpt-4o",
                temperature=0.7,
                messages=[
                    {"role": "system", "content": "You are a talented writer like Stephen King. Write a catchy title and a detailed ebook outline based on the user's idea. Make it engaging and well-structured."},
                    {"role": "user", "content": prompt},
                ]
            )
            return res.choices[0].message.content

    def generate_scenes(outline):
        with st.spinner("✨ Creating vivid scenes..."):
            res = openai_client.chat.completions.create(
                model="gpt-4o",
                temperature=0.8,
                messages=[
                    {"role": "system", "content": "You are a creative novelist. Based on the following ebook outline, generate vivid descriptions of 3-5 key scenes. Each scene should have a clear heading and 2-3 paragraphs of immersive, detailed narrative. Separate each scene with two newlines."},
                    {"role": "user", "content": outline},
                ]
            )
            return res.choices[0].message.content

    def generate_illustration(prompt):
        with st.spinner("🎨 Creating illustration..."):
            img = openai_client.images.generate(
                model="dall-e-3",
                prompt=f"Create a beautiful, artistic illustration for this scene: {prompt}. Style: cinematic, detailed, atmospheric.",
                size="1024x1024"
            )
            return img.data[0].url

    # Main input section
    st.markdown("### 💡 Share Your Ebook Idea")
    st.markdown("*Describe your story concept, theme, or inspiration - just a few sentences are enough!*")
    
    prompt = st.text_area(
        "Your Idea", 
        placeholder="e.g., A mystery novel about a detective who can see the last moments of murder victims, or a guide to sustainable living in urban environments...",
        height=100
    )
    
    if st.button("🚀 Generate Ebook Outline", disabled=not prompt):
        if prompt:
            ebook_text = generate_text(prompt)
            st.session_state["ebook_text"] = ebook_text
            st.session_state["scenes"] = []  # Reset scenes when generating new outline
            st.session_state["images"] = {}   # Reset images

    # Display ebook outline
    if st.session_state["ebook_text"]:
        st.markdown("### 📋 Generated Ebook Outline")
        st.markdown(f"""
        <div class="scene-container">
            {st.session_state["ebook_text"].replace(chr(10), '<br>')}
        </div>
        """, unsafe_allow_html=True)
        
        if st.button("📖 Generate Scene Descriptions"):
            scenes_text = generate_scenes(st.session_state["ebook_text"])
            # Better scene splitting
            scenes = [scene.strip() for scene in scenes_text.split("\n\n") if scene.strip()]
            st.session_state["scenes"] = scenes

    # Display and edit scenes
    if st.session_state["scenes"]:
        st.markdown("### 🎬 Scene Descriptions & Illustrations")
        
        for i, scene in enumerate(st.session_state["scenes"]):
            with st.expander(f"📝 Scene {i+1}", expanded=True):
                col1, col2 = st.columns([2, 1])
                
                with col1:
                    edited_scene = st.text_area(
                        f"Edit Scene {i+1}",
                        scene,
                        height=150,
                        key=f"scene_{i}"
                    )
                    st.session_state["scenes"][i] = edited_scene
                    
                    if st.button(f"🎨 Generate Illustration", key=f"img_btn_{i}"):
                        # Extract first few sentences for better prompt
                        scene_preview = edited_scene[:400] + "..." if len(edited_scene) > 400 else edited_scene
                        image_url = generate_illustration(scene_preview)
                        st.session_state["images"][i] = image_url
                        st.rerun()
                
                with col2:
                    if i in st.session_state["images"]:
                        st.image(
                            st.session_state["images"][i], 
                            caption=f"Scene {i+1} Illustration",
                            use_column_width=True
                        )
                    else:
                        st.markdown("""
                        <div style="background: #f0f2f6; padding: 2rem; border-radius: 10px; text-align: center; color: #666;">
                            🖼️<br>Click "Generate Illustration"<br>to create artwork for this scene
                        </div>
                        """, unsafe_allow_html=True)

        # Export section
        if st.session_state["scenes"]:
            st.markdown("""
            <div class="download-section">
                <h3>📚 Download Your Ebook</h3>
                <p style="color: white; margin-bottom: 1.5rem;">Export your complete ebook with scenes and illustrations</p>
            </div>
            """, unsafe_allow_html=True)
            
            col1, col2, col3 = st.columns([1, 1, 1])
            
            with col2:
                export_format = st.selectbox(
                    "Choose Format",
                    ["DOCX (Recommended)", "PDF (Text Only)"],
                    key="export_format"
                )
                
                if st.button("📥 Download Ebook", type="primary"):
                    if export_format.startswith("DOCX"):
                        with st.spinner("📝 Creating DOCX file..."):
                            buffer = BytesIO()
                            doc = Document()
                            
                            # Add title
                            title_lines = st.session_state["ebook_text"].split('\n')
                            if title_lines:
                                doc.add_heading(title_lines[0].replace('#', '').strip(), 0)
                            
                            # Add outline
                            doc.add_heading("Ebook Outline", level=1)
                            doc.add_paragraph(st.session_state["ebook_text"])
                            doc.add_page_break()
                            
                            # Add scenes with images
                            for i, scene in enumerate(st.session_state["scenes"], start=1):
                                doc.add_heading(f"Scene {i}", level=1)
                                doc.add_paragraph(scene)
                                
                                # Add image if exists
                                if i-1 in st.session_state["images"]:
                                    try:
                                        img_url = st.session_state["images"][i-1]
                                        response = requests.get(img_url)
                                        img_bytes = BytesIO(response.content)
                                        doc.add_picture(img_bytes, width=Inches(5))
                                        doc.add_paragraph("")  # Add spacing
                                    except Exception as e:
                                        doc.add_paragraph(f"[Illustration unavailable: {img_url}]")
                                
                                if i < len(st.session_state["scenes"]):
                                    doc.add_page_break()
                            
                            doc.save(buffer)
                            buffer.seek(0)
                            
                            st.download_button(
                                "📥 Download DOCX",
                                buffer,
                                file_name="my_ebook.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    
                    else:  # PDF
                        with st.spinner("📄 Creating PDF file..."):
                            buffer = BytesIO()
                            pdf = SimpleDocTemplate(buffer)
                            styles = getSampleStyleSheet()
                            story = []
                            
                            # Add title
                            title_lines = st.session_state["ebook_text"].split('\n')
                            if title_lines:
                                story.append(Paragraph(title_lines[0].replace('#', '').strip(), styles["Title"]))
                                story.append(Spacer(1, 24))
                            
                            # Add outline
                            story.append(Paragraph("Ebook Outline", styles["Heading1"]))
                            story.append(Paragraph(st.session_state["ebook_text"].replace("\n", "<br/>"), styles["Normal"]))
                            story.append(Spacer(1, 24))
                            
                            # Add scenes
                            for i, scene in enumerate(st.session_state["scenes"], start=1):
                                story.append(Paragraph(f"Scene {i}", styles["Heading2"]))
                                story.append(Paragraph(scene.replace("\n", "<br/>"), styles["Normal"]))
                                
                                if i-1 in st.session_state["images"]:
                                    story.append(Paragraph(f"Illustration: {st.session_state['images'][i-1]}", styles["Italic"]))
                                
                                story.append(Spacer(1, 20))
                            
                            pdf.build(story)
                            buffer.seek(0)
                            
                            st.download_button(
                                "📥 Download PDF",
                                buffer,
                                file_name="my_ebook.pdf",
                                mime="application/pdf"
                            )

    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 1rem;">
       
    </div>
    """, unsafe_allow_html=True)